# docx_creator.py (финальная версия с заменой сокращений)
import json
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Соответствие сокращений направлений полным названиям
DIRECTION_MAP = {
    "ГПВ": "Мероприятия гражданско-патриотической направленности",
    "ДД": "Мероприятия добровольческой направленности",
    "ДНВ": "Мероприятия в рамках духовно-нравственного воспитания",
    "ЕНН": "Мероприятия естественно-научной направленности",
    "ТН": "Мероприятия творческой направленности",
    "БЖД": "Мероприятия, посвященные безопасности жизнедеятельности и здоровому образу жизни",
    "ЭН": "Мероприятия экологической направленности",
    "РЛК": "Мероприятия, направленные на развитие лидерских качеств",
    "ЦНП": "Мероприятия по воспитанию ценности научного познания"
}

# Словари для замены сокращений на полные названия
FORM_TYPE_MAP = {
    "КлЧ": "Классный час",
    "ВолА": "Волонтерская акция",
    "РодС": "Родительское собрание",
    "През": "Презентация",
    "МК": "Мастер-класс",
    "Конф": "Конференция",
    "Форум": "Форум",
    "Рег": "Региональные и всероссийские акции",
    "Проекты": "Проекты",
    "ИН": "Иные формы",
    "Классный час": "Классный час",
    "Волонтерская акция": "Волонтерская акция",
    "Родительское собрание": "Родительское собрание",
    "Презентация": "Презентация",
    "Мастер-класс": "Мастер-класс",
    "Конференция": "Конференция",
    "Форум": "Форум",
    "Региональные и всероссийские акции": "Региональные и всероссийские акции",
    "Проекты": "Проекты",
    "Иные формы": "Иные формы"
}

PARTICIPATION_FORMAT_MAP = {
    "Оч": "Очный",
    "Дист": "Дистанционный",
    "Смеш": "Смешанный",
    "Очный": "Очный",
    "Дистанционный": "Дистанционный",
    "Смешанный": "Смешанный"
}

def expand_form_type(value):
    """Заменяет сокращение формы проведения на полное название"""
    if not value:
        return value
    
    # Если значение содержит несколько сокращений через запятую или пробел
    original = value.strip()
    
    # Проверяем, нужно ли заменять
    for short, full in FORM_TYPE_MAP.items():
        if original == short:
            return full
    
    # Если значение не является точным сокращением, пробуем найти в строке
    result = original
    for short, full in FORM_TYPE_MAP.items():
        if short in result:
            result = result.replace(short, full)
    
    return result

def expand_participation_format(value):
    """Заменяет сокращение формата участия на полное название"""
    if not value:
        return value
    
    original = value.strip()
    
    for short, full in PARTICIPATION_FORMAT_MAP.items():
        if original == short:
            return full
    
    result = original
    for short, full in PARTICIPATION_FORMAT_MAP.items():
        if short in result:
            result = result.replace(short, full)
    
    return result

def set_cell_background(cell, color):
    """Устанавливает цвет фона ячейки таблицы (например, 'F2F2F2' для светло-серого)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_alignment(cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=None):
    """Устанавливает выравнивание в ячейке по горизонтали и вертикали."""
    if vertical is not None:
        cell.vertical_alignment = vertical
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal

def set_run_font(run, size=12, bold=False, underline=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline

def set_cell_font(cell, size=11, bold=False):
    """Устанавливает шрифт для всего текста в ячейке."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)

def create_document_from_json(input_json_path, output_docx_path, month_name="января", teacher_count="_____"):
    """
    Создаёт документ Word по образцу на основе JSON-файла с мероприятиями.
    :param input_json_path: путь к JSON-файлу (например, doc.json)
    :param output_docx_path: путь для сохранения .docx
    :param month_name: месяц в родительном падеже (например, "февраля")
    :param teacher_count: количество педагогов (по умолчанию прочерк)
    """
    # Загружаем данные
    with open(input_json_path, 'r', encoding='utf-8') as f:
        events = json.load(f)

    # Группируем по направлениям
    grouped = {}
    for ev in events:
        dir_code = ev.get('direction', '').strip()
        if dir_code not in grouped:
            grouped[dir_code] = []
        grouped[dir_code].append(ev)

    # Создаём документ
    doc = Document()

    # --- Устанавливаем альбомную ориентацию и размер A4 ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)

    # Поля: верх 3 см, низ 1.5 см, лево 2 см, право 2 см
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ---- Шапка документа ----
    # Приложение (по правому краю)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Приложение")
    set_run_font(run, size=12, bold=False)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)

    # Объединённая строка: "Организация мероприятий ... в" + подчеркнутая часть (без подчёркивания шрифта)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("Организация мероприятий по основным направлениям воспитательной работы в ")
    set_run_font(run1, size=12, bold=False, underline=False)
    # Линия из символов подчёркивания – не нужно применять underline шрифта
    run2 = p.add_run(f"_________________________ за {month_name} 2026 года")
    set_run_font(run2, size=12, bold=False, underline=False)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)

    # Строка о численности (без отступа сверху)
    p = doc.add_paragraph()
    run = p.add_run("Численность педагогических работников, получающих денежное вознаграждение за классное руководство (кураторство), по состоянию на 13.01.2026 г. - ")
    set_run_font(run, size=12)
    # teacher_count уже содержит символы подчёркивания, просто выводим их
    run = p.add_run(f"{teacher_count} человек")
    set_run_font(run, size=12)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)

    # Небольшой отступ перед таблицей
    doc.add_paragraph()

    # ---- Таблица ----
    headers = [
        "№ п/п",
        "Дата",
        "Название мероприятия",
        "Уровень",
        "Место проведения",
        "Форма проведения мероприятия \n\n(классный час, волонтерская акция, родительское собрание, презентация, мастер-класс, конференции, форумы, региональные и всероссийские акции, проекты, иные)",
        "Формат участия в мероприятии \n\n(очный, дистанционный, смешанный)",
        "Краткое описание, результат",
        "Ссылки на публикации"
    ]

    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.autofit = False

    # Ширина столбцов из оригинала (в сантиметрах)
    cm_widths = [0.75, 2.25, 3.25, 2.75, 2.5, 3.75, 3.25, 5.25, 3.51]
    for i, width_cm in enumerate(cm_widths):
        table.columns[i].width = Cm(width_cm)

    # Заголовок таблицы (светло-серая заливка, шрифт 11, обычный, выравнивание по центру)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], 'F2F2F2')
        set_cell_font(hdr_cells[i], size=11, bold=False)
        set_cell_alignment(hdr_cells[i], horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_PARAGRAPH.CENTER)

    # Проходим по всем направлениям в порядке из DIRECTION_MAP
    ordered_codes = list(DIRECTION_MAP.keys())
    for code in grouped.keys():
        if code not in ordered_codes:
            ordered_codes.append(code)

    for code in ordered_codes:
        # Заголовок направления (объединённая ячейка, жирный, размер 12, по центру)
        row = table.add_row()
        merged_cell = row.cells[0].merge(row.cells[-1])
        direction_title = DIRECTION_MAP.get(code, f"Направление {code}")
        merged_cell.text = direction_title
        for paragraph in merged_cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True)
        set_cell_alignment(merged_cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_PARAGRAPH.CENTER)

        # Добавляем мероприятия, если есть
        if code in grouped and grouped[code]:
            for idx, ev in enumerate(grouped[code], start=1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = ev.get('date', '')
                row.cells[2].text = ev.get('name', '')
                row.cells[3].text = ev.get('level', '')
                row.cells[4].text = ev.get('location', '')
                # Заменяем сокращения на полные названия
                row.cells[5].text = expand_form_type(ev.get('form_type', ''))
                row.cells[6].text = expand_participation_format(ev.get('participation_format', ''))
                row.cells[7].text = ev.get('description_result', '')
                row.cells[8].text = ev.get('link', '')
                for cell in row.cells:
                    set_cell_font(cell, size=11, bold=False)
                    set_cell_alignment(cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            # Пустая строка-разделитель
            row = table.add_row()
            for cell in row.cells:
                cell.text = ""
                set_cell_font(cell, size=11, bold=False)
                set_cell_alignment(cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_PARAGRAPH.CENTER)

        # Пустая строка после группы
        row = table.add_row()
        for cell in row.cells:
            cell.text = ""
            set_cell_font(cell, size=11, bold=False)
            set_cell_alignment(cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_PARAGRAPH.CENTER)

    # Сохраняем документ
    try:
        doc.save(output_docx_path)
        print(f"Документ сохранён: {output_docx_path}")
    except PermissionError:
        base, ext = os.path.splitext(output_docx_path)
        for i in range(1, 100):
            new_path = f"{base}_{i}{ext}"
            if not os.path.exists(new_path):
                try:
                    doc.save(new_path)
                    print(f"Документ сохранён как: {new_path} (файл {output_docx_path} был занят)")
                    return
                except:
                    continue
        print("Не удалось сохранить документ: все имена заняты или недоступны.")

# Пример использования
if __name__ == "__main__":
    create_document_from_json("doc.json", "output.docx", month_name="января", teacher_count="_____")